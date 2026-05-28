from pathlib import Path
from typing import List

from langchain_text_splitters import (
    RecursiveCharacterTextSplitter
)

from langchain_core.documents import Document

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.core.config import settings
from app.utils.logger import get_logger
logger=get_logger()

class DocumentProcessor:
    """
    Handles loading, parsing,
    and chunking of documents.
    """

    def __init__(self):
        self.chunk_size = settings.MAX_CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""]
        )

    def process_document(
        self,
        file_path: str
    ) -> List[Document]:
        """
        Main entry point for document processing.
        """

        file_extension = Path(file_path).suffix.lower()

        logger.info(
            f"Processing document: {file_path}"
        )

        if file_extension == ".pdf":
            text = self._load_pdf(file_path)

        elif file_extension == ".docx":
            text = self._load_docx(file_path)

        elif file_extension == ".txt":
            text = self._load_txt(file_path)

        else:
            raise ValueError(
                f"Unsupported file format: {file_extension}"
            )

        chunks = self._split_text(
            text=text,
            source=file_path
        )

        logger.info(
            f"Created {len(chunks)} chunks"
        )

        return chunks

    def _load_pdf(
        self,
        file_path: str
    ) -> str:
        """
        Load text from PDF.
        """

        text = ""

        reader = PdfReader(file_path)

        for page in reader.pages:
            extracted_text = page.extract_text()

            if extracted_text:
                text += extracted_text + "\n"

        logger.info("PDF loaded successfully")

        return text

    def _load_docx(
        self,
        file_path: str
    ) -> str:
        """
        Load text from DOCX.
        """

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        logger.info("DOCX loaded successfully")

        return text

    def _load_txt(
        self,
        file_path: str
    ) -> str:
        """
        Load text from TXT.
        """

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:
            text = file.read()

        logger.info("TXT loaded successfully")

        return text

    def _split_text(
        self,
        text: str,
        source: str
    ) -> List[Document]:
        """
        Split text into chunks.
        """

        chunks = self.text_splitter.split_text(text)

        documents = []

        for index, chunk in enumerate(chunks):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": source,
                        "chunk_id": index
                    }
                )
            )

        return documents